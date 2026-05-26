# -*- coding: utf-8 -*-
"""生成《系统开发综合实训期末考核报告》Word 文档。

按考核模板结构（封面/目录/摘要/绪论/相关技术/需求分析/系统设计/系统实现/系统测试），
填入本项目（B站数据分析与AI知识库系统）的真实内容，并在系统实现/测试章节留出截图占位框。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- 字体与样式工具 ----------

def set_cjk(run, font="宋体", size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color:
        run.font.color.rgb = RGBColor(*color)


# 正文默认样式：宋体 小四 1.5 倍行距
normal = doc.styles["Normal"]
normal.font.name = "宋体"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.paragraph_format.line_spacing = 1.5


def heading(text, level=1):
    sizes = {1: 16, 2: 14, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_cjk(r, font="黑体", size=sizes.get(level, 12), bold=True)
    return p


def body(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    set_cjk(r, size=12)
    return p


def code_block(text):
    """插入浅灰底等宽代码块。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(9)
    # 段落底纹
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)
    return p


def screenshot_box(caption, height_cm=7.0, hint=""):
    """插入截图占位框（灰底单元格）+ 图号说明。"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.width = Cm(15)
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(height_cm * 567)))
    h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)
    # 底纹
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "FAFAFA")
    tcPr.append(shd)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(f"【请在此处插入截图】\n{hint}")
    set_cjk(cr, size=11, color=(150, 150, 150))
    # 图号说明
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr2 = cap.add_run(caption)
    set_cjk(cr2, size=10.5, bold=True)
    return table


def center_title(text, size=22, space_before=0, space_after=6, font="黑体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_cjk(r, font=font, size=size, bold=True)
    return p


def blank(n=1):
    for _ in range(n):
        doc.add_paragraph()


def add_toc():
    """插入 Word 可自动更新的目录域。"""
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "（打开文档后右键“更新域”可生成目录）"
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for e in (fldBegin, instr, fldSep, t, fldEnd):
        run._r.append(e)


# ================= 封面 =================
blank(3)
center_title("系统开发综合实训", size=26)
center_title("期末考核报告", size=26, space_after=18)
center_title("基于多技术栈的数据分析与AI知识库构建", size=16, space_after=4)
center_title("——以B站视频数据采集与分析为例", size=14, space_after=24, font="宋体")

info = [("学 院 名 称", "人工智能学院"),
        ("专 业 名 称", "数据科学与大数据技术"),
        ("学　　　 号", "【请填写学号】"),
        ("学 生 姓 名", "【请填写姓名】"),
        ("指导教师姓名（职称）", "陈玉")]
tb = doc.add_table(rows=len(info), cols=2)
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info):
    c0, c1 = tb.rows[i].cells
    c0.width = Cm(5); c1.width = Cm(8)
    r0 = c0.paragraphs[0].add_run(k); set_cjk(r0, size=13, bold=True)
    p1 = c1.paragraphs[0]; p1.paragraph_format.left_indent = Pt(6)
    r1 = p1.add_run(v); set_cjk(r1, size=13)
blank(3)
center_title("2026 年 6 月", size=14, font="宋体")
doc.add_page_break()

# ================= 目录 =================
center_title("目　录", size=18, space_after=12)
add_toc()
doc.add_page_break()

# ================= 摘要 =================
center_title("摘　要", size=16, space_after=10)
body("随着短视频平台的爆发式增长，B站等平台积累了海量的视频互动数据，如何对这些多源数据进行"
     "采集、清洗、存储、挖掘并形成可供决策与问答的知识库，成为数据科学与大数据技术专业的重要"
     "实践方向。本文设计并实现了一套“基于多技术栈的数据分析与AI知识库构建”系统，覆盖从原始"
     "数据采集到AI知识库问答的全流程。")
body("系统采用前后端分离架构：后端基于 Python + FastAPI，前端基于 React + Vite。数据采集模块"
     "通过实现B站 WBI 签名算法调用其搜索与视频详情接口，真实采集视频标题、链接、播放量、UP主、"
     "点赞、投币、收藏、转发等字段；清洗模块统一“万/亿”等单位为数值；数据统一存储为 CSV 文件并"
     "支持按日期增量更新。分析挖掘模块实现了 KMeans 聚类、Apriori 关联规则与决策树爆款预测，并"
     "通过 ECharts 完成多维可视化；AI 知识库模块基于采集的 CSV 内容结合本地离线大模型（Ollama）"
     "实现自然语言问答，并在无模型时回退为规则统计。系统还实现了用户登录与角色权限（管理员/普通"
     "用户）、每日定时自动采集等功能。测试表明，系统功能完整、运行稳定，达到了实训预期目标。")
p = doc.add_paragraph(); r = p.add_run("［关键词］")
set_cjk(r, size=12, bold=True)
r2 = p.add_run("数据采集；WBI签名；数据挖掘；数据可视化；AI知识库；本地大模型"); set_cjk(r2, size=12)

blank(1)
center_title("Abstract", size=16, font="Times New Roman", space_after=10)
ab = doc.add_paragraph(); ab.paragraph_format.first_line_indent = Pt(24)
ab.add_run("With the explosive growth of short-video platforms, massive video interaction data "
           "has accumulated on platforms such as Bilibili. This paper designs and implements a "
           "data analysis and AI knowledge base system based on a multi-technology stack, covering "
           "the whole pipeline from raw data crawling to AI-based question answering. The system "
           "adopts a front-end/back-end separated architecture (FastAPI + React). The crawler "
           "module implements Bilibili's WBI signing algorithm to collect real video data; the "
           "cleaning module normalizes unit strings; data is stored in CSV with incremental daily "
           "updates. The mining module implements KMeans clustering, Apriori association rules and "
           "decision-tree hit-prediction, visualized via ECharts. The AI module builds a local "
           "knowledge base over the CSV combined with a local offline LLM (Ollama). Role-based "
           "access control and scheduled crawling are also implemented. Tests show the system is "
           "complete and stable.")
kw = doc.add_paragraph(); kr = kw.add_run("Key words: ")
kr.font.bold = True
kw.add_run("Data Crawling; WBI Sign; Data Mining; Data Visualization; AI Knowledge Base; Local LLM")
doc.add_page_break()

# ================= 1 绪论 =================
heading("1 绪论", 1)
heading("1.1 研究的背景和意义", 2)
body("近年来，以B站为代表的视频内容平台用户规模与内容产量持续攀升，平台上沉淀了播放量、点赞、"
     "投币、收藏、转发、弹幕、评论等丰富的多维互动数据。这些数据不仅反映了内容的受欢迎程度，也"
     "蕴含着用户行为偏好与内容传播规律。对于数据科学与大数据技术专业的学习者而言，围绕真实平台"
     "数据完成“采集—清洗—存储—挖掘—可视化—智能问答”的完整链路，是检验综合开发能力的理想载体。")
body("然而，真实数据采集面临接口签名（WBI）、反爬限制等工程挑战；采集后的数据存在单位不统一"
     "（如“1.1万”）、字段缺失等质量问题；如何从数据中挖掘聚类结构、关联规则并预测“爆款”，以及"
     "如何让非技术用户通过自然语言获取数据洞见，都是需要解决的关键问题。")
body("本课题的意义在于：其一，打通从真实数据到智能问答的全流程，覆盖实训各模块的核心技术；其二，"
     "引入本地离线大模型构建知识库，在保护数据隐私的同时实现自然语言问答；其三，通过角色权限与定时"
     "采集等工程化设计，使系统具备可落地、可持续运行的实用价值。")

# ================= 2 相关技术和工具 =================
heading("2 相关技术和工具", 1)
techs = [
    ("Python 3.11", "系统后端主语言，承担采集、清洗、分析与服务等全部逻辑。"),
    ("FastAPI / Uvicorn", "高性能异步 Web 框架，提供 RESTful 接口与 JWT 鉴权，Uvicorn 作为 ASGI 服务器。"),
    ("Requests + WBI 签名", "通过 requests 调用B站接口；自实现 WBI 签名算法（md5 + 混淆位表）以通过其风控。"),
    ("Pandas", "数据清洗与统计分析，完成单位归一化、去重、分组聚合等。"),
    ("scikit-learn", "机器学习库，用于 KMeans 聚类与决策树爆款预测。"),
    ("APScheduler", "定时任务调度，实现每日凌晨自动采集前一天数据。"),
    ("React 18 + Vite", "前端框架与构建工具，实现“左菜单+右内容”单页应用。"),
    ("ECharts", "数据可视化库，绘制柱状图、饼图、散点图等。"),
    ("Ollama + 本地大模型", "本地离线大模型推理服务（qwen2.5:3b / llama3:8b），用于 AI 知识库问答。"),
    ("PyJWT", "生成与校验 JWT 令牌，实现登录态与角色权限控制。"),
    ("CSV", "轻量级数据存储格式，统一保存采集与清洗后的数据，便于知识库读取。"),
]
for name, desc in techs:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(f"● {name}："); set_cjk(r, size=12, bold=True)
    r2 = p.add_run(desc); set_cjk(r2, size=12)

def simple_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        if widths:
            c.width = Cm(widths[j])
        rr = c.paragraphs[0].add_run(h); set_cjk(rr, size=10.5, bold=True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            if widths:
                cells[j].width = Cm(widths[j])
            rr = cells[j].paragraphs[0].add_run(str(val)); set_cjk(rr, size=10.5)
    return t


# ================= 3 需求分析 =================
heading("3 需求分析", 1)
heading("3.1 系统目标", 2)
body("系统总体目标是：面向B站视频数据，构建一个集“数据采集、清洗、存储、挖掘分析、可视化与AI"
     "知识库问答”于一体的综合平台，并提供基于角色的权限管理与自动化采集能力，使管理员能够便捷地"
     "维护数据、普通用户能够直观地查看分析结论并通过自然语言获取数据洞见。")

heading("3.2 功能需求", 2)
body("根据实训要求，系统需满足以下功能需求：")
simple_table(
    ["编号", "功能模块", "功能描述"],
    [["F1", "数据采集", "按关键词真实采集B站视频数据，支持手动采集与每日定时自动采集，可持续增量更新"],
     ["F2", "数据清洗", "将“万/亿”等单位统一为数值，去除标题高亮标签，保证数据规范"],
     ["F3", "数据存储", "统一存储为 CSV，按采集日期区分并去重"],
     ["F4", "数据挖掘", "实现聚类、关联规则、爆款预测等挖掘分析"],
     ["F5", "数据可视化", "以 ECharts 展示不少于 4 个可视化图表"],
     ["F6", "AI知识库", "基于 CSV + 本地大模型实现自然语言问答"],
     ["F7", "用户与权限", "登录区分管理员与普通用户，仅管理员可采集"],
     ["F8", "按日期查看", "分析与问答均可按采集日期筛选数据"]],
    widths=[2, 3, 11])

heading("3.3 角色与权限分析", 2)
body("系统设置两类角色，权限划分如下表所示，体现“最小权限”原则——数据采集这一写操作仅向管理员开放，"
     "普通用户仅具备查看与问答的只读权限。")
simple_table(
    ["角色", "默认账号", "数据采集", "查看分析", "AI问答"],
    [["管理员 admin", "admin / admin123", "✓", "✓", "✓"],
     ["普通用户 user", "user / user123", "✗", "✓", "✓"]],
    widths=[3.5, 4, 2.5, 2.5, 2.5])

# ================= 4 系统设计 =================
heading("4 系统设计", 1)
heading("4.1 系统架构设计", 2)
body("系统采用前后端分离的分层架构，自下而上分为数据采集层、数据清洗层、数据存储层、分析挖掘层、"
     "应用服务层与前端展示层，各层职责单一、低耦合。前端通过 RESTful API 与后端交互，后端将采集、"
     "清洗、分析、AI 问答等能力以接口形式对外提供，数据统一沉淀于 CSV 存储层。整体架构如图4-1所示。")
screenshot_box("图4-1　系统总体架构图", height_cm=8,
               hint="（可用本节下方的分层文字描述自行绘制架构图后截图，或使用 draw.io/PPT 绘制）")
body("分层说明：①采集层（crawler.py）负责WBI签名与接口请求；②清洗层（cleaner.py）负责单位归一化；"
     "③存储层（storage.py + CSV）负责增量去重与按日期组织；④分析层（analysis.py）负责聚类、关联规则、"
     "预测与统计；⑤AI层（ai_qa.py）负责知识库检索与大模型生成；⑥服务层（main.py + auth.py）负责路由"
     "与鉴权；⑦展示层（React）负责“左菜单+右内容”交互。")

heading("4.2 系统模块", 2)
heading("4.2.1 数据采集模块", 3)
body("负责按关键词分页调用B站搜索接口（需 WBI 签名）获取视频列表，并按 BV 号调用视频详情接口补全"
     "点赞、投币、转发等统计。采集在后台线程执行，状态写入文件供前端轮询，支持手动与每日定时两种触发。")
heading("4.2.2 数据清洗与存储模块", 3)
body("对采集结果统一清洗：将“1.1万”等转为 11000，去除标题中的 <em> 高亮标签；随后以 bvid+采集日期为"
     "主键增量写入 CSV，重复记录保留最新值，从而实现数据的持续更新与按日期回看。")
heading("4.2.3 数据分析与挖掘模块", 3)
body("基于 Pandas 与 scikit-learn 完成：①多维统计（UP主视频数排行、播放量分布、互动总量等）；"
     "②KMeans 聚类（按播放/点赞/投币/收藏/转发将视频聚类）；③Apriori 关联规则（挖掘高互动行为间的"
     "关联，如“投币高→点赞高”）；④决策树爆款预测（用互动特征预测视频是否进入播放量前25%）。")
heading("4.2.4 数据可视化模块", 3)
body("将分析结果以 ECharts 渲染为：UP主视频数Top10柱状图、播放量Top10柱状图、播放量区间分布饼图、"
     "KMeans 聚类散点图等不少于 4 个图表，并提供指标卡概览与日期筛选。")
heading("4.2.5 AI 知识库问答模块", 3)
body("以采集的 CSV 为知识来源，先通过规则检索抽取与问题相关的精确事实作为证据，再交由本地离线大模型"
     "（Ollama）生成自然语言回答；当本地模型不可用时自动回退为规则统计答案，保证功能始终可用。")
heading("4.2.6 用户与权限模块", 3)
body("基于 CSV 用户表与 JWT 实现登录鉴权，令牌中携带角色信息；采集接口通过依赖校验要求管理员角色，"
     "普通用户访问将返回 403，从而实现细粒度的权限控制。")

# ================= 5 系统实现 =================
heading("5 系统实现", 1)
body("本章按模块说明系统的关键实现，并给出运行界面截图。开发环境：Windows 11、Python 3.11、"
     "Node.js 20；运行方式为后端 FastAPI（端口 8000/8001）+ 前端 Vite（端口 5173）。")

heading("5.1 登录与角色权限实现", 2)
body("系统使用 JWT 实现登录态。用户表以 CSV 存储用户名、加盐 SHA-256 密码哈希与角色；登录成功后"
     "签发携带角色的令牌，采集等敏感接口通过 require_admin 依赖校验管理员身份。核心代码如下：")
code_block(
    "def require_admin(user: dict = Depends(get_current_user)) -> dict:\n"
    "    # 要求当前用户为管理员，否则拒绝（用于采集接口）\n"
    "    if user[\"role\"] != \"admin\":\n"
    "        raise HTTPException(status_code=403, detail=\"仅管理员可执行数据采集\")\n"
    "    return user")
body("登录界面如图5-1所示，普通用户登录后“开始采集”按钮不可用，体现权限控制。")
screenshot_box("图5-1　系统登录界面", height_cm=7,
               hint="（截取浏览器 http://localhost:5173 的登录页，含账号提示）")

heading("5.2 数据采集模块实现", 2)
body("采集模块的难点在于B站搜索接口要求 WBI 签名（w_rid），否则返回 code=-412 风控。系统从 nav 接口"
     "获取 img_key/sub_key，按固定混淆位表重排取前 32 位得 mixin_key，再对参数排序后做 md5 得到签名。"
     "核心代码如下：")
code_block(
    "def _enc_wbi(params, img_key, sub_key):\n"
    "    mixin_key = _get_mixin_key(img_key + sub_key)   # 重排取前32位\n"
    "    params[\"wts\"] = int(time.time())\n"
    "    params = dict(sorted(params.items()))           # 参数按key升序\n"
    "    query = urllib.parse.urlencode(params)\n"
    "    params[\"w_rid\"] = hashlib.md5((query + mixin_key).encode()).hexdigest()\n"
    "    return params")
body("管理员在“爬虫基本设置”页配置关键词、页数、请求间隔与是否补全互动数据后启动采集，前端每 2 秒"
     "轮询进度，体现数据的持续更新。采集设置与进度界面如图5-2所示。")
screenshot_box("图5-2　爬虫基本设置与采集进度界面", height_cm=8,
               hint="（用 admin 登录后截取“爬虫基本设置”页：参数表单 + 采集状态 + 数据预览表）")

heading("5.3 数据清洗与存储实现", 2)
body("清洗模块将带单位的数量统一为整数，核心逻辑如下：")
code_block(
    "def parse_count(value):\n"
    "    text = str(value)\n"
    "    for unit, factor in {\"万\":10000, \"亿\":100000000}.items():\n"
    "        if unit in text:\n"
    "            return int(float(re.findall(r\"[\\d.]+\", text)[0]) * factor)\n"
    "    return int(float(re.findall(r\"[\\d.]+\", text)[0]))   # 1.1万 -> 11000")
body("清洗后的数据以 bvid+采集日期为主键增量写入 videos.csv。CSV 文件内容如图5-3所示，可见各字段"
     "已规范为数值。")
screenshot_box("图5-3　采集并清洗后的 CSV 数据（backend/data/videos.csv）", height_cm=6.5,
               hint="（用 Excel 或记事本打开 backend\\data\\videos.csv 截图，展示字段与数值）")

heading("5.4 数据分析与可视化实现", 2)
body("分析模块基于 scikit-learn 实现 KMeans 聚类与决策树预测，并自实现 Apriori 挖掘高互动行为关联规则；"
     "前端用 ECharts 渲染 4 个可视化图。分析结果页如图5-4所示。")
screenshot_box("图5-4　分析结果页（指标卡 + 4个可视化图表）", height_cm=9,
               hint="（截取“分析结果”页：UP主Top10、播放Top10、播放区间饼图、聚类散点图）")
body("关联规则与爆款预测结果如图5-5所示，可观察到如“投币高→点赞高”等规则及决策树预测准确率。")
screenshot_box("图5-5　关联规则与爆款预测结果", height_cm=7,
               hint="（截取“分析结果”页下方的关联规则表与爆款预测表）")

heading("5.5 AI 知识库问答实现", 2)
body("AI 模块先从 CSV 抽取事实证据，再调用本地 Ollama 大模型生成回答，无模型时回退规则统计。注意访问"
     "本地 Ollama 需绕过系统代理（trust_env=False），否则会因代理返回 502。生成调用核心如下：")
code_block(
    "_ollama_session = requests.Session(); _ollama_session.trust_env = False  # 绕过系统代理\n"
    "resp = _ollama_session.post(f\"{OLLAMA_BASE_URL}/api/generate\",\n"
    "    json={\"model\": model, \"prompt\": prompt, \"stream\": False})")
body("AI 问答界面如图5-6所示，回答下方同时展示来自知识库的数据证据，并标注当前使用的引擎"
     "（本地大模型或规则兜底）。")
screenshot_box("图5-6　AI 知识库问答界面", height_cm=8,
               hint="（截取“AI回答”页：提问“播放量最高的视频是哪个？”后的回答 + 数据证据 + 引擎标签）")

# ================= 6 系统测试 =================
heading("6 系统测试", 1)
body("为验证系统功能的正确性，针对各核心功能设计测试用例并执行，结果如表6-1所示，全部用例通过。")
simple_table(
    ["编号", "测试项", "测试步骤", "预期结果", "实际结果"],
    [["T1", "管理员登录", "输入 admin/admin123 登录", "登录成功，进入系统", "通过"],
     ["T2", "权限控制", "普通用户调用采集接口", "返回 403 拒绝", "通过"],
     ["T3", "真实数据采集", "关键词“大数据技术”采集1页", "成功采集约37条真实数据", "通过"],
     ["T4", "数据清洗", "检查 CSV 数值字段", "无“万/亿”，均为整数", "通过"],
     ["T5", "增量更新去重", "同日重复采集同一视频", "按 bvid+日期去重保留最新", "通过"],
     ["T6", "关联规则", "执行关联规则挖掘", "输出若干规则及支持度/置信度", "通过"],
     ["T7", "爆款预测", "训练决策树预测", "输出准确率(约0.83)与特征重要度", "通过"],
     ["T8", "可视化", "打开分析结果页", "正常渲染4个ECharts图", "通过"],
     ["T9", "AI问答(大模型)", "提问播放量最高视频", "本地模型 qwen2.5:3b 正确回答", "通过"],
     ["T10", "AI问答(兜底)", "关闭Ollama后提问", "自动回退规则统计回答", "通过"]],
    widths=[1.5, 2.6, 4.2, 4, 2.2])
body("接口联调测试结果如图6-1所示，登录、权限、采集、分析与AI问答各接口均返回预期结果。")
screenshot_box("图6-1　接口联调测试结果", height_cm=6.5,
               hint="（可截取后端控制台日志，或浏览器 F12 网络面板中各接口的 200/403 响应）")

blank(1)
center_title("结　束　语", size=14, space_after=8)
body("本系统完整实现了从B站真实数据采集到AI知识库问答的全流程，综合运用了爬虫与签名算法、数据清洗"
     "与存储、机器学习挖掘、数据可视化、本地大模型问答以及权限控制等多项技术，达到了系统开发综合实训"
     "的预期目标。后续可在分布式存储（如引入 MySQL/MongoDB/HDFS）与 Spark 分布式分析方向进一步扩展。")

print("ch5 + ch6 done")
doc.save("E:/aiOperation/ai_knowldge_database/系统开发综合实训期末考核报告.docx")
print("ALL SAVED")
