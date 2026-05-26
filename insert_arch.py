# -*- coding: utf-8 -*-
"""把 docs/architecture.png 插入报告的「图4-1」占位框，不改动其它任何内容。
运行前请先关闭 Word。"""
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PATH = "系统开发综合实训期末考核报告.docx"
IMG = "docs/architecture.png"

d = Document(PATH)

target = None
for tbl in d.tables:
    txt = "".join(n.text or "" for n in tbl._tbl.iter(qn("w:t")))
    if "draw.io" in txt or "架构图" in txt:
        target = tbl
        break
if target is None:
    raise SystemExit("未找到图4-1占位表格")

cell = target.cell(0, 0)
for p in cell.paragraphs:          # 清空占位文字
    for r in p.runs:
        r.text = ""
para = cell.paragraphs[0]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.add_run().add_picture(IMG, width=Cm(14.5))
for p in cell.paragraphs[1:]:      # 删除多余空段
    p._element.getparent().remove(p._element)

d.save(PATH)
print("SUCCESS: 架构图已插入图4-1占位框")
