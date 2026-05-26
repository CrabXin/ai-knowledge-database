# -*- coding: utf-8 -*-
"""在现有考核报告上【仅追加】新增功能内容（三级角色 + 用户管理），不改动其它已有内容。
追加点：
  1) 3.3 角色权限表：增加“超级用户”行 + 说明段
  2) 第4章末（“5 系统实现”前）：新增 4.2.7 用户管理模块
  3) 第5章末（“6 系统测试”前）：新增 5.6 用户管理实现 + 图5-7 截图占位
  4) 第6章测试表：追加 T11/T12/T13 三条用例
运行前请先关闭 Word。"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "系统开发综合实训期末考核报告.docx"
doc = Document(PATH)


def set_cjk(run, font="宋体", size=12, bold=False, color=None):
    run.font.size = Pt(size); run.font.bold = bold; run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color:
        run.font.color.rgb = RGBColor(*color)


def find_para(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def find_table(keyword):
    for t in doc.tables:
        if keyword in "".join(n.text or "" for n in t._tbl.iter(qn("w:t"))):
            return t
    return None


# ---- 在某锚点元素前插入的构造器 ----
class Inserter:
    def __init__(self, anchor_el):
        self.anchor = anchor_el

    def _place(self, el):
        self.anchor.addprevious(el)

    def heading(self, text, level=2):
        size = {1: 16, 2: 14, 3: 13}.get(level, 13)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text); set_cjk(r, font="黑体", size=size, bold=True)
        self._place(p._p); return p

    def body(self, text):
        p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24)
        r = p.add_run(text); set_cjk(r, size=12)
        self._place(p._p); return p

    def code(self, text):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(12)
        r = p.add_run(text); r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas"); r.font.size = Pt(9)
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
        p._p.get_or_add_pPr().append(shd)
        self._place(p._p); return p

    def screenshot(self, caption, height_cm=7.5, hint=""):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
        cell = table.cell(0, 0); cell.width = Cm(15)
        tr = table.rows[0]._tr; trPr = tr.get_or_add_trPr()
        h = OxmlElement("w:trHeight"); h.set(qn("w:val"), str(int(height_cm * 567)))
        h.set(qn("w:hRule"), "atLeast"); trPr.append(h)
        tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "FAFAFA"); tcPr.append(shd)
        cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(f"【请在此处插入截图】\n{hint}"); set_cjk(cr, size=11, color=(150, 150, 150))
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr2 = cap.add_run(caption); set_cjk(cr2, size=10.5, bold=True)
        self._place(table._tbl); self._place(cap._p); return table


# ========== 1) 角色表：加“超级用户”行 + 说明段 ==========
role_tbl = find_table("默认账号")
if role_tbl is None:
    raise SystemExit("未找到角色权限表")
ncol = len(role_tbl.columns)
vals = ["超级用户 superadmin", "superadmin / super123", "✓", "✓", "✓"][:ncol]
row = role_tbl.add_row()
for j, v in enumerate(row.cells):
    rr = v.paragraphs[0].add_run(vals[j] if j < len(vals) else "✓"); set_cjk(rr, size=10.5)
# 移动到表头后成为第一数据行
role_tbl.rows[0]._tr.addnext(row._tr)
# 表后追加说明段
note = doc.add_paragraph(); note.paragraph_format.first_line_indent = Pt(24)
nr = note.add_run(
    "在“管理员 / 普通用户”基础上新增“超级用户（superadmin）”角色，形成三级权限层级："
    "超级用户 > 管理员 > 普通用户。其中超级用户可新增管理员或普通用户账号，管理员可新增普通用户账号，"
    "新建账号即时生效、可直接登录并按对应角色获得权限，从而实现分级授权与可扩展的用户管理。")
set_cjk(nr, size=12)
role_tbl._tbl.addnext(note._p)

# ========== 2) “5 系统实现”前：4.2.7 用户管理模块 ==========
anchor5 = find_para("系统实现") or find_para("5 系统实现")
if anchor5 is None:
    raise SystemExit("未找到“系统实现”标题")
ins = Inserter(anchor5._p)
ins.heading("4.2.7 用户管理模块（三级角色权限）", 3)
ins.body("为支持账号的分级管理，系统在登录鉴权基础上扩展了用户管理模块。角色分为超级用户、管理员、"
         "普通用户三级：超级用户负责创建管理员账号，管理员负责创建普通用户账号，实现职责分离与权限下放。"
         "用户信息（用户名、加盐哈希密码、角色）统一存储于 users.csv，新增账号即时生效、可直接登录；"
         "后端通过 can_create_role 校验“谁能创建谁”，越权创建将被拒绝（HTTP 403）。")

# ========== 3) “6 系统测试”前：5.6 用户管理实现 + 图5-7 ==========
anchor6 = find_para("系统测试") or find_para("6 系统测试")
if anchor6 is None:
    raise SystemExit("未找到“系统测试”标题")
ins6 = Inserter(anchor6._p)
ins6.heading("5.6 用户管理（三级角色）实现", 2)
ins6.body("用户管理模块在前端表现为仅对管理员与超级用户可见的“用户管理”菜单，页面提供新增账号表单与"
          "现有账号列表，可选角色由当前登录者角色动态决定（管理员仅可选“普通用户”，超级用户可选"
          "“管理员/普通用户”）。后端创建账号的权限控制核心代码如下：")
ins6.code(
    "def can_create_role(actor_role, target_role):\n"
    "    # 超级用户可建 管理员/普通用户；管理员仅可建 普通用户\n"
    "    if actor_role == \"superadmin\":\n"
    "        return target_role in (\"admin\", \"user\")\n"
    "    if actor_role == \"admin\":\n"
    "        return target_role == \"user\"\n"
    "    return False")
ins6.body("用户管理界面如图5-7所示，超级用户可在此新增管理员或普通用户，新建账号可立即用于登录。")
ins6.screenshot("图5-7　用户管理界面（新增账号与现有账号列表）", height_cm=7.5,
                hint="（用 superadmin 或 admin 登录后截取“用户管理”页：新增账号表单 + 账号列表）")

# ========== 4) 测试表：追加三条用例 ==========
test_tbl = find_table("测试项")
if test_tbl is None:
    raise SystemExit("未找到测试用例表")
cases = [
    ["T11", "超级用户建管理员", "superadmin 新增 admin 账号", "创建成功，可直接登录", "通过"],
    ["T12", "管理员建普通用户", "admin 新增 user 账号", "创建成功，可直接登录", "通过"],
    ["T13", "越权创建拦截", "管理员尝试新增管理员账号", "返回 403 拒绝", "通过"],
]
tcol = len(test_tbl.columns)
for case in cases:
    cells = test_tbl.add_row().cells
    for j in range(tcol):
        rr = cells[j].paragraphs[0].add_run(case[j] if j < len(case) else ""); set_cjk(rr, size=10.5)

doc.save(PATH)
print("SUCCESS: 已在现有报告上追加 三级角色 + 用户管理 相关内容")
