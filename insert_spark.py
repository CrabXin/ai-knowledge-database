# -*- coding: utf-8 -*-
"""在现有报告上【仅追加】模块五（Spark 分布式分析）内容，其它不动。
追加点：
  1) 第2章技术：Redis 条目后补 Spark(PySpark) 一条
  2) 第4章末（系统实现前）：新增 4.2.8 Spark 分布式分析模块
  3) 第5章末（系统测试前）：新增 5.8 Spark 分布式分析实现 + 图5-9 截图占位
  4) 功能需求表：追加 F10 分布式分析
  5) 测试表：追加 T15 Spark 分布式分析
运行前请先关闭 Word。"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "系统开发综合实训期末考核报告.docx"
doc = Document(PATH)


def set_cjk(run, font="宋体", size=12, bold=False, color=None):
    run.font.size = Pt(size); run.font.bold = bold; run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color:
        run.font.color.rgb = RGBColor(*color)


def find_para(pred):
    for p in doc.paragraphs:
        if pred(p.text.strip()):
            return p
    return None


def find_table(keyword):
    for t in doc.tables:
        if keyword in "".join(n.text or "" for n in t._tbl.iter(qn("w:t"))):
            return t
    return None


class Inserter:
    def __init__(self, anchor_el):
        self.anchor = anchor_el

    def heading(self, text, level=2):
        size = {1: 16, 2: 14, 3: 13}.get(level, 13)
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text); set_cjk(r, font="黑体", size=size, bold=True)
        self.anchor.addprevious(p._p); return p

    def body(self, text):
        p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24)
        r = p.add_run(text); set_cjk(r, size=12)
        self.anchor.addprevious(p._p); return p

    def code(self, text):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(12)
        r = p.add_run(text); r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas"); r.font.size = Pt(9)
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
        p._p.get_or_add_pPr().append(shd)
        self.anchor.addprevious(p._p); return p

    def screenshot(self, caption, height_cm=7.0, hint=""):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
        cell = table.cell(0, 0); cell.width = Cm(15)
        trPr = table.rows[0]._tr.get_or_add_trPr()
        h = OxmlElement("w:trHeight"); h.set(qn("w:val"), str(int(height_cm * 567))); h.set(qn("w:hRule"), "atLeast"); trPr.append(h)
        tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "FAFAFA"); tcPr.append(shd)
        cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(f"【请在此处插入截图】\n{hint}"); set_cjk(cr, size=11, color=(150, 150, 150))
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr2 = cap.add_run(caption); set_cjk(cr2, size=10.5, bold=True)
        self.anchor.addprevious(table._tbl); self.anchor.addprevious(cap._p); return table


done = []

# 1) 第2章：Redis 条目后补 Spark
redis_p = find_para(lambda t: t.startswith("●") and "Redis" in t and "键值" in t)
if redis_p is None:
    redis_p = find_para(lambda t: "Redis" in t and "排行榜" in t)
if redis_p is not None:
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run("● Spark（PySpark）："); set_cjk(r, size=12, bold=True)
    r2 = p.add_run("分布式计算引擎。以本地 SparkSession 对数据做分布式聚合分析，"
                   "结合 Spark SQL 与 DataFrame 算子完成多维统计。"); set_cjk(r2, size=12)
    redis_p._p.addnext(p._p)
    done.append("第2章技术: +Spark")

# 2) 第4章末（系统实现前）：4.2.8 Spark 分布式分析模块
anchor_impl = find_para(lambda t: t in ("系统实现", "5 系统实现"))
if anchor_impl is not None:
    ins = Inserter(anchor_impl._p)
    ins.heading("4.2.8 Spark 分布式分析模块", 3)
    ins.body("为满足实训模块五“使用 Spark 进行分布式数据分析”的要求，系统在单机分析之外引入 Spark。"
             "通过本地 SparkSession（local[*]）将清洗后的数据加载为 Spark DataFrame，"
             "结合 Spark SQL 与 DataFrame 算子完成不少于 4 种分布式分析：UP主视频数排行、播放量统计、"
             "播放量区间分布、各互动指标均值、视频时长分桶与平均播放量的关系。Spark 不可用时自动回退"
             "为 pandas 等价实现，保证分析结果始终可得。")
    done.append("4.2.8: +Spark模块设计")

# 3) 第5章末（系统测试前）：5.8 Spark 实现 + 图5-9
anchor_test = find_para(lambda t: t in ("系统测试", "6 系统测试"))
if anchor_test is not None:
    ins = Inserter(anchor_test._p)
    ins.heading("5.8 Spark 分布式分析实现", 2)
    ins.body("Spark 分析通过 /api/analysis/spark 接口触发，后端懒加载本地 SparkSession（复用），"
             "将 pandas 数据转为 Spark DataFrame 并注册临时视图，再用 Spark SQL 与 DataFrame 算子计算。"
             "Windows 环境下需绑定 127.0.0.1 并指定 PYSPARK_PYTHON 以解决 Python worker 回连问题，"
             "Java 17 需追加 add-opens 参数。核心代码如下：")
    ins.code(
        "spark = (SparkSession.builder.master(\"local[*]\")\n"
        "         .config(\"spark.driver.host\", \"127.0.0.1\")\n"
        "         .config(\"spark.driver.bindAddress\", \"127.0.0.1\").getOrCreate())\n"
        "sdf = spark.createDataFrame(pdf); sdf.createOrReplaceTempView(\"videos\")\n"
        "# 分布式分析①：UP主视频数 Top10（Spark SQL）\n"
        "spark.sql(\"SELECT author, COUNT(*) n FROM videos GROUP BY author \"\n"
        "          \"ORDER BY n DESC LIMIT 10\").collect()")
    ins.body("前端“分析结果”页底部提供“运行 Spark 分布式分析”按钮，计算完成后以图表展示 5 种分析结果，"
             "并标注当前计算引擎（Spark 版本）。运行结果如图5-9所示。")
    ins.screenshot("图5-9　Spark 分布式分析结果（含引擎标识与多维图表）", height_cm=8,
                   hint="（截取“分析结果”页底部 Spark 区：引擎标签 Spark 3.5.1 + UP主排行/区间分布/时长分桶/互动均值等图表）")
    done.append("第5章: +5.8 Spark实现 + 图5-9")

# 4) 功能需求表：追加 F10
freq = find_table("功能模块")
if freq is not None:
    cells = freq.add_row().cells
    vals = ["F10", "分布式分析", "使用 Spark 对数据进行不少于 4 种分布式聚合分析"]
    for j in range(len(freq.columns)):
        rr = cells[j].paragraphs[0].add_run(vals[j] if j < len(vals) else ""); set_cjk(rr, size=10.5)
    done.append("功能需求表: +F10")

# 5) 测试表：追加 T15
testtbl = find_table("测试项")
if testtbl is not None:
    cells = testtbl.add_row().cells
    vals = ["T15", "Spark分布式分析", "点击运行 Spark 分析", "返回5种分析结果，引擎为 Spark 3.5.1", "通过"]
    for j in range(len(testtbl.columns)):
        rr = cells[j].paragraphs[0].add_run(vals[j] if j < len(vals) else ""); set_cjk(rr, size=10.5)
    done.append("测试表: +T15")

doc.save(PATH)
print("SUCCESS 追加完成:")
for d in done:
    print("  -", d)
