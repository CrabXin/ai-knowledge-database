# -*- coding: utf-8 -*-
"""在现有报告上【仅追加】模块三（多存储系统 MongoDB + Redis）相关内容，其它不动。
追加点：
  1) 第2章 相关技术：CSV 条目后补 MongoDB / Redis 两条
  2) 4.2.2 末尾：补一段“双存储系统”设计说明
  3) 第5章末（系统测试前）：新增 5.7 多存储系统实现 + 图5-8 截图占位
  4) 功能需求表：追加 F9 多存储入库
  5) 测试表：追加 T14 多存储入库
运行前请先关闭 Word。"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "系统开发综合实训期末考核报告.docx"
doc = Document(PATH)


def set_cjk(run, font="宋体", size=12, bold=False, color=None):
    run.font.size = Pt(size); run.font.bold = bold; run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color:
        run.font.color.rgb = RGBColor(*color)


def find_para(pred):
    for p in doc.paragraphs:
        if pred(p.text.strip()):
            return p
    return None


def find_table(keyword):
    for t in doc.tables:
        if keyword in "".join(n.text or "" for n in t._tbl.iter(qn("w:t"))):
            return t
    return None


def make_bullet(name, desc):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(f"● {name}："); set_cjk(r, size=12, bold=True)
    r2 = p.add_run(desc); set_cjk(r2, size=12)
    return p


class Inserter:
    def __init__(self, anchor_el):
        self.anchor = anchor_el

    def heading(self, text, level=2):
        size = {1: 16, 2: 14, 3: 13}.get(level, 13)
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text); set_cjk(r, font="黑体", size=size, bold=True)
        self.anchor.addprevious(p._p); return p

    def body(self, text):
        p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24)
        r = p.add_run(text); set_cjk(r, size=12)
        self.anchor.addprevious(p._p); return p

    def code(self, text):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(12)
        r = p.add_run(text); r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas"); r.font.size = Pt(9)
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
        p._p.get_or_add_pPr().append(shd)
        self.anchor.addprevious(p._p); return p

    def screenshot(self, caption, height_cm=7.0, hint=""):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
        cell = table.cell(0, 0); cell.width = Cm(15)
        trPr = table.rows[0]._tr.get_or_add_trPr()
        h = OxmlElement("w:trHeight"); h.set(qn("w:val"), str(int(height_cm * 567))); h.set(qn("w:hRule"), "atLeast"); trPr.append(h)
        tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "FAFAFA"); tcPr.append(shd)
        cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(f"【请在此处插入截图】\n{hint}"); set_cjk(cr, size=11, color=(150, 150, 150))
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr2 = cap.add_run(caption); set_cjk(cr2, size=10.5, bold=True)
        self.anchor.addprevious(table._tbl); self.anchor.addprevious(cap._p); return table


done = []

# 1) 第2章：CSV 条目后补 MongoDB / Redis
csv_p = find_para(lambda t: "CSV" in t and "存储格式" in t)
if csv_p is not None:
    b_redis = make_bullet("Redis", "内存型键值数据库。为每条视频建立 Hash，用 Set 维护当日索引、用 Sorted Set 维护播放量排行榜，支持高速读取与 Top-N 查询。")
    b_mongo = make_bullet("MongoDB", "面向文档的 NoSQL 数据库。以文档形式存储每条视频，按 (bvid, 采集日期) upsert，支持灵活查询与聚合分析。")
    csv_p._p.addnext(b_mongo._p)      # 先放 Mongo 紧跟 CSV
    b_mongo._p.addnext(b_redis._p)    # Redis 再跟其后
    done.append("第2章技术: +MongoDB/Redis")

# 2) 4.2.2 末尾（4.2.3 前）：双存储设计说明
anchor_423 = find_para(lambda t: t.startswith("4.2.3"))
if anchor_423 is not None:
    Inserter(anchor_423._p).body(
        "在 CSV 主存储之外，系统按实训模块三“任选 2 个以上存储系统”的要求，将清洗后的数据同步写入 "
        "MongoDB 与 Redis 两个存储系统：MongoDB 以 (bvid, 采集日期) 为唯一键执行 upsert，并建立索引便于按"
        "日期、播放量查询；Redis 为每条视频建立 Hash，并以集合维护当日索引、以有序集合维护播放量排行榜，"
        "便于 Top-N 等快速查询。任一存储系统不可用时自动跳过，不影响 CSV 主流程与系统运行。")
    done.append("4.2.2: +双存储说明")

# 3) 第5章末（系统测试前）：5.7 多存储系统实现 + 图5-8
anchor_test = find_para(lambda t: t in ("系统测试", "6 系统测试"))
if anchor_test is not None:
    ins = Inserter(anchor_test._p)
    ins.heading("5.7 多存储系统（MongoDB + Redis）实现", 2)
    ins.body("模块三要求将清洗后的数据存入 MySQL / MongoDB / Redis / HDFS 中的两个以上，本系统选用 "
             "MongoDB + Redis。采集流程在清洗后调用统一入库函数 save_all 同步写入两库，并提供 "
             "/api/storage/status 接口返回各库的连接状态与数据量。核心代码如下：")
    ins.code(
        "# MongoDB：按 (bvid, crawl_date) upsert，避免重复\n"
        "col.update_one({\"bvid\": rec[\"bvid\"], \"crawl_date\": rec[\"crawl_date\"]},\n"
        "               {\"$set\": rec}, upsert=True)\n\n"
        "# Redis：Hash 存每条视频 + 有序集合维护播放量排行榜\n"
        "r.hset(f\"video:{date}:{bvid}\", mapping=rec)\n"
        "r.zadd(f\"rank:play:{date}\", {bvid: rec[\"play\"]})")
    ins.body("前端“爬虫基本设置”页实时展示两个存储系统的连接状态与数据量，如图5-8所示，"
             "可见 MongoDB 与 Redis 均已连接且数据量与采集结果一致。")
    ins.screenshot("图5-8　存储系统连接状态（MongoDB + Redis）", height_cm=6.5,
                   hint="（截取“爬虫基本设置”页的“存储系统状态”卡片：两个绿色卡片显示已连接版本与数据量）")
    done.append("第5章: +5.7 多存储实现 + 图5-8")

# 4) 功能需求表：追加 F9
freq = find_table("功能模块")
if freq is not None:
    cells = freq.add_row().cells
    vals = ["F9", "多存储入库", "清洗后数据同步写入 MongoDB + Redis 两个存储系统"]
    for j in range(len(freq.columns)):
        rr = cells[j].paragraphs[0].add_run(vals[j] if j < len(vals) else ""); set_cjk(rr, size=10.5)
    done.append("功能需求表: +F9")

# 5) 测试表：追加 T14
testtbl = find_table("测试项")
if testtbl is not None:
    cells = testtbl.add_row().cells
    vals = ["T14", "多存储入库", "采集后查看 /api/storage/status", "MongoDB 与 Redis 均连接且数据量一致", "通过"]
    for j in range(len(testtbl.columns)):
        rr = cells[j].paragraphs[0].add_run(vals[j] if j < len(vals) else ""); set_cjk(rr, size=10.5)
    done.append("测试表: +T14")

doc.save(PATH)
print("SUCCESS 追加完成:")
for d in done:
    print("  -", d)
